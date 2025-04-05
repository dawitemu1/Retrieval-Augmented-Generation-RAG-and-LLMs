import os
import streamlit as st
import torch
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
import docx  # python-docx
from transformers import GPT2Tokenizer, GPT2Model, pipeline
from langchain.chains import RetrievalQA
from langchain_community.vectorstores import FAISS
from langchain.embeddings.base import Embeddings
from langchain.schema import Document
from langchain_community.llms import HuggingFacePipeline
import numpy as np

# Prevent Transformers from using TensorFlow
os.environ["TRANSFORMERS_NO_TF"] = "1"

# FAISS database path
FAISS_DB_PATH = "faiss_db"

# Load GPT-2 tokenizer and model
tokenizer = GPT2Tokenizer.from_pretrained("gpt2")
model = GPT2Model.from_pretrained("gpt2")
tokenizer.pad_token = tokenizer.eos_token

# Define custom embedding class
class GPT2Embedding(Embeddings):
    def __init__(self, model, tokenizer):
        self.model = model
        self.tokenizer = tokenizer

    def embed_documents(self, texts):
        embeddings = []
        for text in texts:
            inputs = self.tokenizer(text, return_tensors="pt", truncation=True, padding=True)
            outputs = self.model(**inputs)
            embeddings.append(outputs.last_hidden_state.mean(dim=1).squeeze().detach().numpy())
        return embeddings

    def embed_query(self, text):
        inputs = self.tokenizer(text, return_tensors="pt", truncation=True, padding=True)
        outputs = self.model(**inputs)
        return outputs.last_hidden_state.mean(dim=1).squeeze().detach().numpy()

# Initialize embedding model
gpt2_embedding = GPT2Embedding(model, tokenizer)

# Function to extract text from uploaded files
def extract_text(file):
    file_type = file.name.lower()

    if file_type.endswith(".pdf"):
        text = ""
        doc = fitz.open(stream=file.read(), filetype="pdf")
        for page in doc:
            text += page.get_text()
        return text

    elif file_type.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

    elif file_type.endswith((".png", ".jpg", ".jpeg")):
        image = Image.open(file)
        return pytesseract.image_to_string(image)

    else:
        return None

# Sidebar file upload
st.sidebar.title("📤 Upload Files here")
uploaded_files = st.sidebar.file_uploader("Upload PDF, DOCX, or Image files", type=["pdf", "docx", "png", "jpg", "jpeg"], accept_multiple_files=True)

if uploaded_files:
    new_docs = []
    for file in uploaded_files:
        extracted_text = extract_text(file)
        if extracted_text:
            new_docs.append(Document(page_content=extracted_text, metadata={"source": file.name}))
        else:
            st.sidebar.warning(f"Couldn't extract text from {file.name}")

    if new_docs:
        st.sidebar.success(f"{len(new_docs)} documents processed and added to FAISS!")

        # Embed and save to FAISS
        texts = [doc.page_content for doc in new_docs]
        metadatas = [doc.metadata for doc in new_docs]
        embeddings = gpt2_embedding.embed_documents(texts)

        # Correctly structure the text_embeddings as a list of tuples (text, embedding)
        text_embeddings = list(zip(texts, embeddings))  # This should be a list of (text, embedding) tuples

        # Load or create FAISS DB
        if os.path.exists(FAISS_DB_PATH):
            vector_store = FAISS.load_local(FAISS_DB_PATH, gpt2_embedding, allow_dangerous_deserialization=True)
            vector_store.add_embeddings(text_embeddings)  # Pass the list of tuples directly to FAISS
        else:
            vector_store = FAISS.from_embeddings(embeddings, metadatas, gpt2_embedding)

        vector_store.save_local(FAISS_DB_PATH)

# Load FAISS vector store for QA
if os.path.exists(FAISS_DB_PATH):
    try:
        vector_store = FAISS.load_local(FAISS_DB_PATH, gpt2_embedding, allow_dangerous_deserialization=True)
        retriever = vector_store.as_retriever()
    except Exception as e:
        st.error(f"Error loading FAISS database: {e}")
        st.stop()
else:
    st.error("FAISS database not found. Please upload and process documents first.")
    st.stop()

# Set up the HuggingFace pipeline
generator = pipeline("text-generation", model="gpt2", max_new_tokens=100)
llm = HuggingFacePipeline(pipeline=generator)

# Create QA chain
qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)

# Streamlit Chat UI
st.title("📄 Bank Procedure & Guideline Assistant")
st.write("Ask questions about the Commercial Bank of Ethiopia's procedures and guidelines.")

# Chat state
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "question_count" not in st.session_state:
    st.session_state.question_count = 0

# Display chat history
for msg in st.session_state.chat_history:
    with st.chat_message("user"):
        st.markdown(msg["user"])
    with st.chat_message("assistant"):
        st.markdown(msg["bot"])

# Input box
user_message = st.chat_input("Type your question and press Enter to Get Answer...")

if user_message:
    st.chat_message("user").markdown(user_message)

    # Custom prompt
    prompt = (
        "You are a helpful assistant trained on Commercial Bank of Ethiopia's procedures and guidelines.\n"
        "Please provide concise and clear answers. Answer the following:\n\n"
        f"User: {user_message.strip()}\nAssistant:"
    )

    try:
        # Run QA
        response = qa_chain.run(prompt)

        # Get the source of the document(s)
        sources = []
        similarities = []

        if retriever:
            results = retriever.get_relevant_documents(user_message)
            if results:
                # Calculate cosine similarity for each result
                query_embedding = gpt2_embedding.embed_query(user_message)
                for result in results:
                    # Ensure embeddings are available in the result
                    if hasattr(result, "embedding"):
                        doc_embedding = np.array(result.embedding)
                        similarity = np.dot(query_embedding, doc_embedding) / (np.linalg.norm(query_embedding) * np.linalg.norm(doc_embedding))
                        similarities.append((result.metadata.get('source', 'Unknown document'), similarity))

                # Sort results by similarity (descending)
                similarities.sort(key=lambda x: x[1], reverse=True)

                # Add top relevant sources
                sources = [f"🔹 Source: {source} (Similarity: {similarity:.4f})" for source, similarity in similarities]
            else:
                sources.append("No relevant information found. Please modify your query or provide more context.")

        # Show response with sources in a Google-like search results style
        with st.chat_message("assistant"):
            st.write(f"💡 **Answer:** {response}")
            st.write(" Document Retrieval Results (Ranked by Similarity):")
            for idx, source in enumerate(sources, 1):
                st.write(f"{idx}. {source}")

        # Save chat
        st.session_state.chat_history.append({
            "user": user_message.strip(),
            "bot": f"💡 **Answer:** {response}\n" + "\n".join(sources)
        })

        if st.session_state.question_count == 0:
            st.info("I'm here to help more if you have additional questions!")

        st.session_state.question_count += 1

    except Exception as e:
        st.error(f"Error generating response: {e}")
